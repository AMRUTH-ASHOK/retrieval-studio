"""
Data Type Definitions and Handlers
Supports multiple data types with their own processing logic
"""
from abc import ABC, abstractmethod
from typing import List, Dict, Any, Optional
from dataclasses import dataclass
import uuid


@dataclass
class Document:
    """Standard document representation"""
    doc_id: str
    doc_name: str
    text: str
    metadata: Dict[str, Any]
    data_type: str


class DataTypeHandler(ABC):
    """Base class for data type handlers"""
    
    @abstractmethod
    def get_name(self) -> str:
        """Return the name of this data type"""
        pass
    
    @abstractmethod
    def get_display_name(self) -> str:
        """Return user-friendly display name"""
        pass
    
    @abstractmethod
    def get_input_schema(self) -> Dict:
        """Return the input schema for this data type"""
        pass
    
    @abstractmethod
    def load_documents(self, config: Dict[str, Any]) -> List[Document]:
        """Load and convert data into standard Document format"""
        pass
    
    @abstractmethod
    def get_compatible_strategies(self) -> List[str]:
        """Return list of compatible chunking strategies for this data type"""
        pass


class PDFDataType(DataTypeHandler):
    """PDF document handler"""
    
    def get_name(self) -> str:
        return "pdf"
    
    def get_display_name(self) -> str:
        return "PDF Documents"
    
    def get_input_schema(self) -> Dict:
        return {
            "source_type": "upload",
            "file_types": ["pdf"],
            "multiple": True,
            "fields": [
                {"name": "extract_images", "type": "bool", "default": False, "label": "Extract Images"},
                {"name": "ocr_enabled", "type": "bool", "default": False, "label": "Enable OCR"}
            ]
        }
    
    def load_documents(self, config: Dict[str, Any]) -> List[Document]:
        """Load PDF documents"""
        import fitz  # PyMuPDF

        documents = []
        uploaded_files = config.get("uploaded_files", [])
        extract_images = config.get("extract_images", False)
        ocr_enabled = config.get("ocr_enabled", False)

        for file in uploaded_files:
            filename = getattr(file, 'name', 'unknown.pdf')

            # Extract text from PDF
            text = self._extract_text_from_pdf(file, ocr_enabled)

            # Get page count for metadata
            try:
                # Re-read file for page count (file was already read in _extract_text_from_pdf)
                if hasattr(file, 'seek'):
                    file.seek(0)  # Reset file pointer if possible

                content = file.read() if hasattr(file, 'read') else file
                if isinstance(content, str):
                    content = content.encode('utf-8')

                pdf_doc = fitz.open(stream=content, filetype="pdf")
                page_count = pdf_doc.page_count
                pdf_doc.close()
            except Exception:
                page_count = 0

            doc = Document(
                doc_id=str(uuid.uuid4()),
                doc_name=filename,
                text=text,
                metadata={
                    "source_type": "pdf",
                    "filename": filename,
                    "page_count": page_count,
                    "extract_images": extract_images,
                    "ocr_enabled": ocr_enabled,
                    "file_size": getattr(file, 'size', 0)
                },
                data_type="pdf"
            )
            documents.append(doc)

        return documents
    
    def _extract_text_from_pdf(self, file, ocr_enabled: bool) -> str:
        """Extract text from PDF file using PyMuPDF"""
        try:
            import fitz  # PyMuPDF

            # Read file content
            content = file.read() if hasattr(file, 'read') else file
            if isinstance(content, str):
                content = content.encode('utf-8')

            # Open PDF from bytes
            pdf_document = fitz.open(stream=content, filetype="pdf")

            # Extract text from all pages
            text_parts = []
            for page_num in range(pdf_document.page_count):
                page = pdf_document[page_num]
                page_text = page.get_text()

                if page_text.strip():
                    text_parts.append(page_text)

            pdf_document.close()

            # Combine all pages with page separators
            full_text = "\n\n--- Page Break ---\n\n".join(text_parts)

            if not full_text.strip():
                # If no text extracted, might be a scanned PDF
                if ocr_enabled:
                    return f"[OCR not yet implemented for {getattr(file, 'name', 'file')}]"
                else:
                    return f"[No text found in {getattr(file, 'name', 'file')} - may be scanned. Enable OCR if needed.]"

            return full_text

        except Exception as e:
            # Handle errors gracefully
            filename = getattr(file, 'name', 'file')
            return f"[Error extracting PDF content from {filename}: {str(e)}]"
    
    def get_compatible_strategies(self) -> List[str]:
        return ["baseline", "structured", "parent_child", "semantic"]


class DeltaTableDataType(DataTypeHandler):
    """Delta Table handler"""
    
    def get_name(self) -> str:
        return "delta_table"
    
    def get_display_name(self) -> str:
        return "Delta Table"
    
    def get_input_schema(self) -> Dict:
        return {
            "source_type": "text_input",
            "fields": [
                {"name": "table_name", "type": "text", "required": True, "label": "Table Name (catalog.schema.table)"},
                {"name": "text_column", "type": "text", "required": True, "label": "Text Column Name", "default": "text"},
                {"name": "id_column", "type": "text", "required": False, "label": "ID Column Name (optional)"},
                {"name": "filter_condition", "type": "text", "required": False, "label": "WHERE clause (optional)"}
            ]
        }
    
    def load_documents(self, config: Dict[str, Any]) -> List[Document]:
        """Load documents from Delta Table(s)

        Supports two formats:
        1. Single table: {"table_name": "...", "text_column": "...", ...}
        2. Multiple tables: {"tables": [{"table_name": "...", "text_column": "...", ...}, ...]}
        """
        from utils.query_builder import escape_identifier, sanitize_string

        sql_connector = config.get("sql_connector")

        if not sql_connector:
            raise ValueError("SQL connector required for Delta Table data type")

        documents = []

        # Check if using new array format
        tables = config.get("tables", [])

        if tables:
            # New format: multiple tables
            for table_config in tables:
                table_docs = self._load_from_table(sql_connector, table_config)
                documents.extend(table_docs)
        else:
            # Old format: single table (backward compatibility)
            table_docs = self._load_from_table(sql_connector, config)
            documents.extend(table_docs)

        return documents

    def _load_from_table(self, sql_connector, config: Dict[str, Any]) -> List[Document]:
        """Load documents from a single Delta Table"""
        from utils.query_builder import escape_identifier, sanitize_string

        table_name = config.get("table_name", "")
        text_column = config.get("text_column", "text")
        id_column = config.get("id_column")
        filter_condition = config.get("filter_condition")

        if not table_name:
            return []

        # Build query
        table_parts = table_name.split(".")
        if len(table_parts) == 3:
            catalog, schema, table = table_parts
            full_table = f"{escape_identifier(catalog)}.{escape_identifier(schema)}.{escape_identifier(table)}"
        else:
            full_table = escape_identifier(table_name)

        id_col = escape_identifier(id_column) if id_column else "uuid()"
        text_col = escape_identifier(text_column)

        query = f"SELECT {id_col} as id, {text_col} as text FROM {full_table}"
        if filter_condition:
            query += f" WHERE {filter_condition}"

        results = sql_connector.execute(query)

        documents = []
        for row in results:
            doc = Document(
                doc_id=str(row.get('id', uuid.uuid4())),
                doc_name=f"{table_name}_row_{row.get('id', uuid.uuid4())}",
                text=str(row.get('text', '')),
                metadata={"source_table": table_name, "source_type": "delta_table"},
                data_type="delta_table"
            )
            documents.append(doc)

        return documents
    
    def get_compatible_strategies(self) -> List[str]:
        return ["baseline", "semantic", "sentence"]


class UCVolumeDataType(DataTypeHandler):
    """Unity Catalog Volume handler"""
    
    def get_name(self) -> str:
        return "uc_volume"
    
    def get_display_name(self) -> str:
        return "UC Volume (Files)"
    
    def get_input_schema(self) -> Dict:
        return {
            "source_type": "text_input",
            "fields": [
                {"name": "volume_path", "type": "text", "required": True, "label": "Volume Path"},
                {"name": "file_pattern", "type": "text", "default": "*.txt", "label": "File Pattern"},
                {"name": "recursive", "type": "bool", "default": False, "label": "Search Recursively"}
            ]
        }
    
    def load_documents(self, config: Dict[str, Any]) -> List[Document]:
        """Load documents from UC Volume"""
        import fnmatch

        volume_path = config.get("volume_path", "")
        file_pattern = config.get("file_pattern", "*.txt")
        recursive = config.get("recursive", False)

        if not volume_path:
            raise ValueError("volume_path is required for UC Volume data type")

        documents = []

        try:
            # Import dbutils for Databricks file system access
            from databricks.sdk.runtime import dbutils

            # List files in the volume
            files_to_process = []

            if recursive:
                files_to_process = self._list_files_recursive(dbutils, volume_path, file_pattern)
            else:
                try:
                    all_files = dbutils.fs.ls(volume_path)
                    print(f"[DEBUG UCVolume] Listed {len(all_files)} items in '{volume_path}'")
                    files_to_process = [
                        f.path for f in all_files
                        if not f.isDir() and fnmatch.fnmatch(f.name, file_pattern)
                    ]
                    print(f"[DEBUG UCVolume] {len(files_to_process)} files match pattern '{file_pattern}'")
                except Exception as e:
                    raise ValueError(f"Failed to list files in volume path '{volume_path}': {str(e)}")

            # Process each file
            for file_path in files_to_process:
                try:
                    docs = self._load_file_from_volume(dbutils, file_path, config)
                    if docs:
                        if isinstance(docs, list):
                            documents.extend(docs)
                        else:
                            documents.append(docs)
                        print(f"[DEBUG UCVolume] Loaded {len(docs) if isinstance(docs, list) else 1} doc(s) from {file_path}")
                    else:
                        print(f"[WARNING UCVolume] _load_file_from_volume returned empty for {file_path}")
                except Exception as e:
                    print(f"[WARNING] Failed to load file {file_path}: {str(e)}")
                    continue

            print(f"[INFO UCVolume] Total documents loaded: {len(documents)}")
            return documents

        except ImportError:
            raise ValueError("dbutils not available - UC Volume access requires Databricks runtime")
        except Exception as e:
            raise ValueError(f"Failed to load documents from UC Volume: {str(e)}")

    def _list_files_recursive(self, dbutils, path: str, pattern: str) -> List[str]:
        """Recursively list files matching pattern"""
        import fnmatch

        files = []

        try:
            items = dbutils.fs.ls(path)

            for item in items:
                if item.isDir():
                    # Recursively process subdirectories
                    files.extend(self._list_files_recursive(dbutils, item.path, pattern))
                else:
                    # Check if file matches pattern
                    if fnmatch.fnmatch(item.name, pattern):
                        files.append(item.path)

        except Exception as e:
            print(f"[WARNING] Failed to list directory {path}: {str(e)}")

        return files

    def _load_file_from_volume(self, dbutils, file_path: str, config: Dict[str, Any]) -> List[Document]:
        """Load a single file from UC Volume and route to appropriate handler"""
        import os

        filename = os.path.basename(file_path)
        file_ext = os.path.splitext(filename)[1].lower()

        # Read file content (limit to 10MB)
        try:
            content = dbutils.fs.head(file_path, maxBytes=10*1024*1024)
        except Exception as e:
            print(f"[ERROR] Failed to read file {file_path}: {str(e)}")
            return []

        # Route based on file extension
        if file_ext == '.txt':
            # Plain text file
            doc = Document(
                doc_id=str(uuid.uuid4()),
                doc_name=filename,
                text=content,
                metadata={
                    "source_type": "uc_volume",
                    "file_path": file_path,
                    "file_extension": file_ext
                },
                data_type="uc_volume"
            )
            return [doc]

        elif file_ext == '.pdf':
            # PDF file - use PDF handler
            pdf_handler = PDFDataType()
            # Create a file-like object from content
            import io
            file_obj = io.BytesIO(content.encode('utf-8') if isinstance(content, str) else content)
            file_obj.name = filename

            pdf_docs = pdf_handler.load_documents({
                "uploaded_files": [file_obj],
                "extract_images": config.get("extract_images", False),
                "ocr_enabled": config.get("ocr_enabled", False)
            })

            # Update metadata to include UC Volume source
            for doc in pdf_docs:
                doc.metadata["source_type"] = "uc_volume"
                doc.metadata["file_path"] = file_path
                doc.data_type = "uc_volume"

            return pdf_docs

        elif file_ext == '.csv':
            # CSV file - use CSV handler
            csv_handler = CSVDataType()
            import io
            file_obj = io.StringIO(content if isinstance(content, str) else content.decode('utf-8'))
            file_obj.name = filename

            csv_docs = csv_handler.load_documents({
                "uploaded_files": [file_obj],
                "text_column": config.get("text_column", "text"),
                "id_column": config.get("id_column"),
                "has_header": config.get("has_header", True)
            })

            # Update metadata
            for doc in csv_docs:
                doc.metadata["source_type"] = "uc_volume"
                doc.metadata["file_path"] = file_path
                doc.data_type = "uc_volume"

            return csv_docs

        elif file_ext == '.json':
            # JSON file - use JSON handler
            json_handler = JSONDataType()
            import io
            file_obj = io.StringIO(content if isinstance(content, str) else content.decode('utf-8'))
            file_obj.name = filename

            json_docs = json_handler.load_documents({
                "uploaded_files": [file_obj],
                "text_field": config.get("text_field", "text"),
                "id_field": config.get("id_field"),
                "is_array": config.get("is_array", True)
            })

            # Update metadata
            for doc in json_docs:
                doc.metadata["source_type"] = "uc_volume"
                doc.metadata["file_path"] = file_path
                doc.data_type = "uc_volume"

            return json_docs

        elif file_ext in ('.docx', '.doc'):
            docx_handler = DocxDataType()
            import io
            file_obj = io.BytesIO(content.encode('utf-8') if isinstance(content, str) else content)
            file_obj.name = filename

            docx_docs = docx_handler.load_documents({
                "uploaded_files": [file_obj]
            })

            for doc in docx_docs:
                doc.metadata["source_type"] = "uc_volume"
                doc.metadata["file_path"] = file_path
                doc.data_type = "uc_volume"

            return docx_docs

        else:
            print(f"[WARNING] Unsupported file type: {file_ext} for file {file_path}")
            return []
    
    def get_compatible_strategies(self) -> List[str]:
        return ["baseline", "structured", "parent_child", "semantic"]


class DocxDataType(DataTypeHandler):
    """Word/DOCX document handler"""

    def get_name(self) -> str:
        return "docx"

    def get_display_name(self) -> str:
        return "Word Documents"

    def get_input_schema(self) -> Dict:
        return {
            "source_type": "upload",
            "file_types": ["docx"],
            "multiple": True,
            "fields": []
        }

    def load_documents(self, config: Dict[str, Any]) -> List[Document]:
        """Load Word/DOCX documents using python-docx"""
        documents = []
        uploaded_files = config.get("uploaded_files", [])

        for file in uploaded_files:
            filename = getattr(file, 'name', 'unknown.docx')
            text = self._extract_text_from_docx(file)

            doc = Document(
                doc_id=str(uuid.uuid4()),
                doc_name=filename,
                text=text,
                metadata={
                    "source_type": "docx",
                    "filename": filename,
                    "file_size": getattr(file, 'size', 0)
                },
                data_type="docx"
            )
            documents.append(doc)

        return documents

    def _extract_text_from_docx(self, file) -> str:
        """Extract text from DOCX file preserving paragraphs, headings, and table content"""
        try:
            from docx import Document as DocxDocument
            import io

            content = file.read() if hasattr(file, 'read') else file
            if isinstance(content, str):
                content = content.encode('utf-8')

            doc = DocxDocument(io.BytesIO(content))

            text_parts = []

            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    if paragraph.style and paragraph.style.name.startswith('Heading'):
                        text_parts.append(f"\n{'#' * int(paragraph.style.name[-1]) if paragraph.style.name[-1].isdigit() else '##'} {paragraph.text}\n")
                    else:
                        text_parts.append(paragraph.text)

            for table in doc.tables:
                table_rows = []
                for row in table.rows:
                    cells = [cell.text.strip() for cell in row.cells]
                    table_rows.append(" | ".join(cells))
                if table_rows:
                    text_parts.append("\n" + "\n".join(table_rows) + "\n")

            return "\n\n".join(text_parts) if text_parts else f"[No text found in {getattr(file, 'name', 'file')}]"

        except Exception as e:
            filename = getattr(file, 'name', 'file')
            return f"[Error extracting DOCX content from {filename}: {str(e)}]"

    def get_compatible_strategies(self) -> List[str]:
        return ["baseline", "structured", "parent_child", "semantic"]


class TextDataType(DataTypeHandler):
    """Plain text input handler"""
    
    def get_name(self) -> str:
        return "text"
    
    def get_display_name(self) -> str:
        return "Plain Text"
    
    def get_input_schema(self) -> Dict:
        return {
            "source_type": "text_area",
            "fields": [
                {"name": "text_content", "type": "textarea", "required": True, "label": "Enter or Paste Text"},
                {"name": "document_name", "type": "text", "default": "text_document", "label": "Document Name"}
            ]
        }
    
    def load_documents(self, config: Dict[str, Any]) -> List[Document]:
        """Load plain text document(s)

        Supports two formats:
        1. Single text: {"text_content": "...", "document_name": "..."}
        2. Multiple texts: {"text_entries": [{"text_content": "...", "document_name": "..."}, ...]}
        """
        documents = []

        # Check if using new array format
        text_entries = config.get("text_entries", [])

        if text_entries:
            # New format: multiple text entries
            for idx, entry in enumerate(text_entries):
                text_content = entry.get("text_content", "")
                doc_name = entry.get("document_name", f"text_document_{idx + 1}")

                if text_content.strip():
                    doc = Document(
                        doc_id=str(uuid.uuid4()),
                        doc_name=doc_name,
                        text=text_content.strip(),
                        metadata={"source_type": "text", "entry_index": idx},
                        data_type="text"
                    )
                    documents.append(doc)
        else:
            # Old format: single text entry (backward compatibility)
            text_content = config.get("text_content", "")
            doc_name = config.get("document_name", "text_document")

            if text_content.strip():
                doc = Document(
                    doc_id=str(uuid.uuid4()),
                    doc_name=doc_name,
                    text=text_content.strip(),
                    metadata={"source_type": "text"},
                    data_type="text"
                )
                documents.append(doc)

        return documents
    
    def get_compatible_strategies(self) -> List[str]:
        return ["baseline", "semantic", "sentence", "paragraph"]


class CSVDataType(DataTypeHandler):
    """CSV file handler"""
    
    def get_name(self) -> str:
        return "csv"
    
    def get_display_name(self) -> str:
        return "CSV Files"
    
    def get_input_schema(self) -> Dict:
        return {
            "source_type": "upload",
            "file_types": ["csv"],
            "multiple": True,
            "fields": [
                {"name": "text_column", "type": "text", "required": True, "label": "Text Column Name"},
                {"name": "id_column", "type": "text", "required": False, "label": "ID Column (optional)"},
                {"name": "has_header", "type": "bool", "default": True, "label": "Has Header Row"}
            ]
        }
    
    def load_documents(self, config: Dict[str, Any]) -> List[Document]:
        """Load documents from CSV"""
        import csv
        import io
        
        uploaded_files = config.get("uploaded_files", [])
        text_column = config.get("text_column", "")
        id_column = config.get("id_column")
        has_header = config.get("has_header", True)
        
        documents = []
        
        for file in uploaded_files:
            content = file.read() if hasattr(file, 'read') else str(file)
            if isinstance(content, bytes):
                content = content.decode('utf-8')
            
            reader = csv.DictReader(io.StringIO(content)) if has_header else csv.reader(io.StringIO(content))
            
            for idx, row in enumerate(reader):
                if has_header:
                    text = row.get(text_column, "")
                    doc_id = row.get(id_column, str(uuid.uuid4())) if id_column else str(uuid.uuid4())
                else:
                    text = str(row)
                    doc_id = str(uuid.uuid4())
                
                if text:
                    doc = Document(
                        doc_id=str(doc_id),
                        doc_name=f"{getattr(file, 'name', 'csv')}_row_{idx}",
                        text=text,
                        metadata={"source_file": getattr(file, 'name', 'csv'), "row_index": idx},
                        data_type="csv"
                    )
                    documents.append(doc)
        
        return documents
    
    def get_compatible_strategies(self) -> List[str]:
        return ["baseline", "semantic", "sentence"]


class JSONDataType(DataTypeHandler):
    """JSON file handler"""
    
    def get_name(self) -> str:
        return "json"
    
    def get_display_name(self) -> str:
        return "JSON Files"
    
    def get_input_schema(self) -> Dict:
        return {
            "source_type": "upload",
            "file_types": ["json"],
            "multiple": True,
            "fields": [
                {"name": "text_field", "type": "text", "required": True, "label": "Text Field Path (e.g., 'content' or 'data.text')"},
                {"name": "id_field", "type": "text", "required": False, "label": "ID Field Path (optional)"},
                {"name": "is_array", "type": "bool", "default": True, "label": "JSON is an Array"}
            ]
        }
    
    def load_documents(self, config: Dict[str, Any]) -> List[Document]:
        """Load documents from JSON"""
        import json
        
        uploaded_files = config.get("uploaded_files", [])
        text_field = config.get("text_field", "")
        id_field = config.get("id_field")
        is_array = config.get("is_array", True)
        
        documents = []
        
        for file in uploaded_files:
            content = file.read() if hasattr(file, 'read') else str(file)
            if isinstance(content, bytes):
                content = content.decode('utf-8')
            
            data = json.loads(content)
            items = data if is_array else [data]
            
            for idx, item in enumerate(items):
                text = self._get_nested_field(item, text_field)
                doc_id = self._get_nested_field(item, id_field) if id_field else str(uuid.uuid4())
                
                if text:
                    doc = Document(
                        doc_id=str(doc_id),
                        doc_name=f"{getattr(file, 'name', 'json')}_{idx}",
                        text=str(text),
                        metadata={"source_file": getattr(file, 'name', 'json'), "item_index": idx},
                        data_type="json"
                    )
                    documents.append(doc)
        
        return documents
    
    def _get_nested_field(self, obj: Dict, field_path: str) -> Any:
        """Get nested field from object using dot notation"""
        if not field_path:
            return None
        
        parts = field_path.split('.')
        current = obj
        
        for part in parts:
            if isinstance(current, dict):
                current = current.get(part)
            else:
                return None
        
        return current
    
    def get_compatible_strategies(self) -> List[str]:
        return ["baseline", "semantic", "sentence", "structured"]


# Registry of all available data types
DATA_TYPE_REGISTRY = {
    "pdf": PDFDataType(),
    "docx": DocxDataType(),
    "delta_table": DeltaTableDataType(),
    "uc_volume": UCVolumeDataType(),
    "text": TextDataType(),
    "csv": CSVDataType(),
    "json": JSONDataType()
}


def get_data_type_handler(data_type: str) -> DataTypeHandler:
    """Get handler for a specific data type"""
    handler = DATA_TYPE_REGISTRY.get(data_type)
    if not handler:
        raise ValueError(f"Unknown data type: {data_type}")
    return handler


def get_all_data_types() -> List[DataTypeHandler]:
    """Get all available data type handlers"""
    return list(DATA_TYPE_REGISTRY.values())
